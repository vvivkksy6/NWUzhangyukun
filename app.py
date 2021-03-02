import os
from flask import Flask,request,json,send_from_directory,make_response, jsonify
from flask import render_template
from datetime import timedelta
from docx import Document
import matplotlib.pyplot as plt
from docx.shared import Inches

app = Flask(__name__)
app.config['SEND_FILE_MAX_AGE_DEFAULT'] = timedelta(seconds=1)


@app.route('/')
def index():
    return render_template("Main.html")

@app.route('/', methods=['GET', 'POST'])
def generate_file():
    '''
    负责生成对应文本
    :return:
    '''
    if request.method == 'POST':
        data = request.get_data()
        dict = json.loads(data)

        # 基本信息-------------------------------------------------------------------------------------------------------
        document = Document()
        document.add_heading('疫情报告', 0)
        document.add_heading('基本数据', 1)
        if_cluster = '聚集关闭'
        if_hos = '医院关闭'
        hos_speed = '1档（最慢速度）'
        is_mask = '否'

        if dict['info']['mask']:
            is_mask = '是'
        else:
            is_mask = '否'

        if dict['info']['cluster'] == '2':
            if_cluster = '聚集开启'
        else:
            if_cluster = '聚集关闭'

        if dict['info']['hos']:
            if_hos = '医院开启'
        else:
            if_hos = '医院关闭'

        if dict['info']['hosSpeed'] == 3000:
            hos_speed = '1档（最慢速度）'
        elif dict['info']['hosSpeed'] == 2000:
            hos_speed = '2档（中等速度）'
        elif dict['info']['hosSpeed'] == 1000:
            hos_speed = '1档（最快速度）'

        hos_info = ''
        if if_hos == '医院开启':
            hos_info = '医院容纳数量：' + str(dict['info']['hoscap']) + '\n' + '医院收纳速度：' + hos_speed + '\n'

        document.add_paragraph('初始感染人数：' + str(dict['info']['infectNum']) + '\n' +
                               '初始人群数量：' + str(dict['info']['peopleNum']) + '\n' +
                               '人群聚集状态：' + if_cluster + '\n' +
                               '人群移动速度：' + str(dict['info']['peopleSpeed']) + '档（共五档）\n' +
                               '是否佩戴口罩：' + is_mask + '\n' +
                               '是否开启医院：' + if_hos + '\n' + hos_info)

        plt.rcParams['font.sans-serif'] = ['SimHei']
        plt.rcParams['axes.unicode_minus'] = False
        labels = '初始感染人数', '初始人群数量'
        sizes = [dict['info']['infectNum'], dict['info']['peopleNum']]
        explode = (0, 0.1)
        fig1, ax1 = plt.subplots()
        ax1.pie(sizes, explode=explode, labels=labels, autopct='%1.1f%%',
                shadow=True, startangle=90)
        ax1.axis('equal')
        plt.savefig('basic_info.jpg')
        document.add_picture('basic_info.jpg', width=Inches(6.15))

        document.add_paragraph('根据国家统计局以及相关论文的数据统计，模型参数设置为：\n'\
                               '单个小球模拟为一个人，小球直径为30px长度，当两人距离为两个身位60px或者发生碰撞时，此时可能发生感染事件。'\
                               '其中，蓝色代表潜伏期，红色代表已经感染，棕色不移动代表死亡，橙色代表此时已康复拥有抗体，黑色代表健康未感染。'\
                               '\n·当未戴口罩时：潜伏期有抗体1%感染几率。潜伏期无抗体5%感染几率。感染期有抗体5%感染几率，感染期无抗体30%感染几率。'\
                               '\n·当佩戴口罩时：潜伏期有抗体0.3%感染几率。潜伏期无抗体1%感染几率。感染期有抗体3%感染几率，感染期无抗体15感染几率。\n'\
                               '人群可能发生聚集，当小球之间长时间距离过短时候，发生感染的概率自然会大幅度上升。人群移动为随机移动，一共设置为5档。'\
                               '医院开启之后，设置收容速度为3档（慢中快），每次收容数量为10人，医院容纳数量可自主设置。')

        # 当前数据图------------------------------------------------------------------------------------------------------
        plt.cla()
        plt.title('当前时间图', fontsize=20)
        document.add_heading('数据分析', 1)
        document.add_paragraph('红色-当前感染人数 蓝色-当前潜伏人数 绿色-当前死亡人数\n反应此时横坐标时间点的疫情情况')
        x = []
        for i in range(0, len(dict['count_infect'])):
            x.append(i)
        # 已存在感染人数纵坐标
        y1 = dict['count_infect']
        # 已存在潜伏者人数
        y2 = dict['count_latent']
        # 已存在死亡人数
        y3 = dict['count_die']
        plt.plot(x, y1, 's-', color='r', label="当前感染人数")  # s-:方形
        plt.plot(x, y2, 'o-', color='b', label="当前潜伏人数")  # o-:圆形
        plt.plot(x, y3, 'x-', color='g', label='当前死亡人数')  # x-:x形
        plt.xlabel("时间/每0.1秒")
        plt.ylabel("人数")
        plt.savefig('current_info.jpg')
        document.add_picture('current_info.jpg', width=Inches(5.75))

        # 累计数据图------------------------------------------------------------------------------------------------------
        plt.cla()
        plt.title('累计时间图', fontsize=20)
        document.add_paragraph('红色-累计感染人数 蓝色-累计潜伏人数 绿色-累计死亡人数\n反应此时横坐标时间点累计一共的疫情情况')
        x = []
        for i in range(0, len(dict['count_infect'])):
            x.append(i)
        # 已存在感染人数纵坐标
        y1 = dict['ac_infect']
        # 已存在潜伏者人数
        y2 = dict['ac_latent']
        # 已存在死亡人数
        y3 = dict['ac_die']
        plt.plot(x, y1, 's-', color='r', label="累计感染人数")  # s-:方形
        plt.plot(x, y2, 'o-', color='b', label="累计潜伏人数")  # o-:圆形
        plt.plot(x, y3, 'x-', color='g', label='累计死亡人数')  # x-:x形
        plt.xlabel("时间/每0.1秒")
        plt.ylabel("人数")
        plt.savefig('ac_info.jpg')
        document.add_picture('ac_info.jpg', width=Inches(5.75))

        # 死亡数据图------------------------------------------------------------------------------------------------------
        document.add_paragraph('\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n')
        document.add_paragraph('总人数：' + str(dict['info']['infectNum'] + dict['info']['peopleNum']) + '\n'\
                               '死亡人数：' + str(dict['dieNum']))
        plt.cla()
        plt.title('死亡占比图', fontsize=20)
        labels = '死亡人数', '总人数'
        sizes = [dict['dieNum'], dict['info']['infectNum'] + dict['info']['peopleNum']]
        explode = (0, 0.1)
        fig1, ax1 = plt.subplots()
        ax1.pie(sizes, explode=explode, labels=labels, autopct='%1.1f%%',
                shadow=True, startangle=90)
        ax1.axis('equal')
        plt.savefig('die_info.jpg')
        document.add_picture('die_info.jpg', width=Inches(4.95))

        # 感染占比图------------------------------------------------------------------------------------------------------
        document.add_paragraph('\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n')
        document.add_paragraph('总人数：' + str(dict['info']['infectNum'] + dict['info']['peopleNum']) + '\n'\
                               '感染人数：' + str(dict['count_infect_before']))
        plt.cla()
        plt.title('感染占比图', fontsize=20)
        labels = '感染人数', '总人数'
        sizes = [dict['count_infect_before'], dict['info']['infectNum'] + dict['info']['peopleNum']]
        explode = (0, 0.1)
        fig1, ax1 = plt.subplots()
        ax1.pie(sizes, explode=explode, labels=labels, autopct='%1.1f%%',
                shadow=True, startangle=90)
        ax1.axis('equal')
        plt.savefig('infect_info.jpg')
        document.add_picture('infect_info.jpg', width=Inches(4.95))

        # 数理模型-------------------------------------------------------------------------------------------------------
        document.add_paragraph('\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n')
        document.add_heading('数学模型', 1)
        document.add_picture('0001.jpg', width=Inches(6.15))
        document.add_picture('0002.jpg', width=Inches(6.15))
        document.add_picture('0003.jpg', width=Inches(6.15))
        document.save('报告.docx')

        return jsonify(msg="success")
    else:
        return '<h1>只接受post请求！<h1>'

@app.route('/file/', methods=['GET', 'POST'])
def get_file():
    dirpath = os.path.join(app.root_path, "./")
    response = make_response(
        send_from_directory(dirpath, "报告.docx", as_attachment = True))
    return response

if __name__ == '__main__':
    # app.run(host='127.0.0.1', port=5000)
    app.run(host='0.0.0.0', port=8080)
